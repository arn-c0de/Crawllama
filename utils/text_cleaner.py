def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extract text from a PDF file using PyPDF2.
    """
    try:
        from PyPDF2 import PdfReader # type: ignore
    except ImportError:
        raise RuntimeError("PyPDF2 not installed. Please add to requirements.txt.")
    reader = PdfReader(pdf_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text

def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract text from a DOCX file using python-docx.
    """
    try:
        import docx # type: ignore
    except ImportError:
        raise RuntimeError("python-docx not installed. Please add to requirements.txt.")
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])
"""Text cleaning and processing utilities."""
import re
import logging
from typing import Optional
from bs4 import BeautifulSoup

logger = logging.getLogger("crawllama")

# Fallback for when tiktoken is not available
CHARS_PER_TOKEN = 4

# Try to import tiktoken for accurate token counting
try:
    import tiktoken
    TIKTOKEN_AVAILABLE = True
except ImportError:
    TIKTOKEN_AVAILABLE = False
    logger.debug("tiktoken not available, using approximate token counting")


class TextCleaner:
    """Unified text cleaning and processing with both char and token-based operations."""
    
    def __init__(self, model_name: str = "gpt-3.5-turbo"):
        """
        Initialize text cleaner.
        
        Args:
            model_name: Model name for tokenizer selection (for token-based operations)
        """
        self.model_name = model_name
        
        # Initialize tokenizer if tiktoken is available
        if TIKTOKEN_AVAILABLE:
            try:
                self.encoding = tiktoken.encoding_for_model(model_name)
                logger.debug(f"TextCleaner initialized with tiktoken: model={model_name}")
            except KeyError:
                self.encoding = tiktoken.get_encoding("cl100k_base")
                logger.debug("TextCleaner initialized with cl100k_base encoding")
        else:
            self.encoding = None
    
    def estimate_tokens(self, text: str) -> int:
        """
        Estimate token count for text.
        
        Args:
            text: Input text
            
        Returns:
            Token count (accurate if tiktoken available, approximate otherwise)
        """
        if not text:
            return 0
        
        if self.encoding is not None:
            try:
                return len(self.encoding.encode(text))
            except Exception as e:
                logger.warning(f"tiktoken encoding failed: {e}")
                return len(text) // CHARS_PER_TOKEN
        else:
            return len(text) // CHARS_PER_TOKEN
    
    def truncate_by_tokens(self, text: str, max_tokens: int, ellipsis: str = "...") -> str:
        """
        Truncate text to fit within token limit (token-aware).
        
        Args:
            text: Text to truncate
            max_tokens: Maximum tokens allowed
            ellipsis: String to append if truncated
            
        Returns:
            Truncated text
        """
        if not text:
            return text
        
        estimated_tokens = self.estimate_tokens(text)
        
        if estimated_tokens <= max_tokens:
            return text
        
        # Use tiktoken for accurate truncation if available
        if self.encoding is not None:
            try:
                tokens = self.encoding.encode(text)
                truncated_tokens = tokens[:max_tokens]
                truncated = self.encoding.decode(truncated_tokens)
                logger.debug(f"Truncated by tokens: {estimated_tokens} -> {len(truncated_tokens)} tokens")
                return truncated + ellipsis
            except Exception as e:
                logger.warning(f"tiktoken truncation failed: {e}")
                # Fall through to approximate method
        
        # Fallback: approximate truncation
        max_chars = max_tokens * CHARS_PER_TOKEN
        truncated = text[:max_chars]
        last_space = truncated.rfind(' ')
        
        if last_space > max_chars // 2:
            truncated = truncated[:last_space]
        
        logger.debug(f"Truncated by tokens (approximate): {estimated_tokens} -> ~{max_tokens} tokens")
        return truncated + ellipsis
    
    def truncate_by_chars(self, text: str, max_chars: int, ellipsis: str = "...") -> str:
        """
        Truncate text to maximum character length, breaking at word boundaries.
        
        Args:
            text: Text to truncate
            max_chars: Maximum number of characters
            ellipsis: String to append if truncated
            
        Returns:
            Truncated text
        """
        if len(text) <= max_chars:
            return text
        
        # Find last space before max_chars
        truncated = text[:max_chars]
        last_space = truncated.rfind(' ')
        
        if last_space > 0:
            truncated = truncated[:last_space]
        
        return truncated + ellipsis
    
    def clean_html(self, html: str, max_length: Optional[int] = 8000) -> str:
        """
        Extract clean text from HTML, removing scripts and styles.
        
        Args:
            html: Raw HTML content
            max_length: Maximum length of output text (default 8000)
            
        Returns:
            Cleaned text content
        """
        soup = BeautifulSoup(html, "html5lib")
        
        # Remove scripts, styles, nav, aside
        for tag in soup(["script", "style", "nav", "aside"]):
            tag.decompose()
        
        # Extract relevant text elements
        texts = []
        for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", 
                                    "article", "footer", "header", "address", "div"]):
            text = tag.get_text(strip=True)
            if len(text) > 15:  # Only meaningful text
                texts.append(text)
        
        # Combine texts
        combined = "\n".join(texts)
        
        # Clean up whitespace
        combined = re.sub(r'\s+', ' ', combined)
        combined = combined.strip()
        
        # Truncate if needed
        if max_length and len(combined) > max_length:
            combined = combined[:max_length] + "..."
        
        return combined
    
    def clean_whitespace(self, text: str) -> str:
        """
        Normalize whitespace in text.
        
        Args:
            text: Text to clean
            
        Returns:
            Cleaned text
        """
        # Replace multiple spaces with single space
        text = re.sub(r' +', ' ', text)
        
        # Replace multiple newlines with double newline
        text = re.sub(r'\n\s*\n+', '\n\n', text)
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def extract_contact_info(self, html: str) -> dict:
        """
        Extract contact information from HTML.
        
        Args:
            html: Raw HTML content
            
        Returns:
            Dictionary with contact info (emails, phones, addresses)
        """
        soup = BeautifulSoup(html, "html5lib")
        text = soup.get_text()
        
        contact_info = {
            "emails": [],
            "phones": [],
            "addresses": []
        }
        
        # Extract emails
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        contact_info["emails"] = list(set(emails))
        
        # Extract phone numbers
        phone_patterns = [
            r'\+?\d{1,4}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}',
            r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\d{4,5}[-.\s]?\d{4,}'
        ]
        phones = []
        for pattern in phone_patterns:
            phones.extend(re.findall(pattern, text))
        contact_info["phones"] = [p for p in set(phones) if len(re.sub(r'\D', '', p)) >= 6]
        
        return contact_info
    
    def extract_urls(self, text: str) -> list:
        """
        Extract URLs from text.
        
        Args:
            text: Text containing URLs
            
        Returns:
            List of extracted URLs
        """
        url_pattern = r'https?://[^\s<>"{}|\\^`\[\]]+'
        return re.findall(url_pattern, text)
    
    def remove_urls(self, text: str) -> str:
        """
        Remove URLs from text.
        
        Args:
            text: Text containing URLs
            
        Returns:
            Text with URLs removed
        """
        url_pattern = r'https?://[^\s<>"{}|\\^`\[\]]+'
        return re.sub(url_pattern, '', text)


# Global instance for convenient access
_text_cleaner_instance = None

def get_text_cleaner(model_name: str = "gpt-3.5-turbo") -> TextCleaner:
    """
    Get or create global TextCleaner instance.
    
    Args:
        model_name: Model name for tokenizer
        
    Returns:
        TextCleaner instance
    """
    global _text_cleaner_instance
    if _text_cleaner_instance is None:
        _text_cleaner_instance = TextCleaner(model_name)
    return _text_cleaner_instance


# ============================================================================
# DEPRECATED: Legacy standalone functions (for backwards compatibility)
# ============================================================================


def extract_contact_info(html: str) -> dict:
    """
    DEPRECATED: Use TextCleaner.extract_contact_info() instead.
    Extract contact information from HTML.
    """
    import warnings
    warnings.warn(
        "extract_contact_info() is deprecated. Use get_text_cleaner().extract_contact_info() instead.",
        DeprecationWarning,
        stacklevel=2
    )
    return get_text_cleaner().extract_contact_info(html)


def clean_html(html: str, max_length: Optional[int] = 8000) -> str:
    """
    DEPRECATED: Use TextCleaner.clean_html() instead.
    Extract clean text from HTML.
    """
    import warnings
    warnings.warn(
        "clean_html() is deprecated. Use get_text_cleaner().clean_html() instead.",
        DeprecationWarning,
        stacklevel=2
    )
    return get_text_cleaner().clean_html(html, max_length)


def truncate_text(text: str, max_chars: int = 1000, ellipsis: str = "...") -> str:
    """
    DEPRECATED: Use TextCleaner.truncate_by_chars() instead.
    Truncate text to maximum length, breaking at word boundaries.
    """
    import warnings
    warnings.warn(
        "truncate_text() is deprecated. Use get_text_cleaner().truncate_by_chars() instead.",
        DeprecationWarning,
        stacklevel=2
    )
    return get_text_cleaner().truncate_by_chars(text, max_chars, ellipsis)


def clean_whitespace(text: str) -> str:
    """
    DEPRECATED: Use TextCleaner.clean_whitespace() instead.
    Normalize whitespace in text.
    """
    import warnings
    warnings.warn(
        "clean_whitespace() is deprecated. Use get_text_cleaner().clean_whitespace() instead.",
        DeprecationWarning,
        stacklevel=2
    )
    return get_text_cleaner().clean_whitespace(text)


def extract_urls(text: str) -> list:
    """
    DEPRECATED: Use TextCleaner.extract_urls() instead.
    Extract URLs from text.
    """
    import warnings
    warnings.warn(
        "extract_urls() is deprecated. Use get_text_cleaner().extract_urls() instead.",
        DeprecationWarning,
        stacklevel=2
    )
    return get_text_cleaner().extract_urls(text)


def remove_urls(text: str) -> str:
    """
    DEPRECATED: Use TextCleaner.remove_urls() instead.
    Remove URLs from text.
    """
    import warnings
    warnings.warn(
        "remove_urls() is deprecated. Use get_text_cleaner().remove_urls() instead.",
        DeprecationWarning,
        stacklevel=2
    )
    return get_text_cleaner().remove_urls(text)


# ============================================================================
# PDF and DOCX extraction (kept as standalone functions)
# ============================================================================


def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extract text from a PDF file using PyPDF2.
    
    Args:
        pdf_path: Path to PDF file
        
    Returns:
        Extracted text content
    """
    try:
        from PyPDF2 import PdfReader  # type: ignore
    except ImportError:
        raise RuntimeError("PyPDF2 not installed. Please add to requirements.txt.")
    reader = PdfReader(pdf_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract text from a DOCX file using python-docx.
    
    Args:
        docx_path: Path to DOCX file
        
    Returns:
        Extracted text content
    """
    try:
        import docx  # type: ignore
    except ImportError:
        raise RuntimeError("python-docx not installed. Please add to requirements.txt.")
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])
